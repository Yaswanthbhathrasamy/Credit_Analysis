"""
CAM Report Generation Service
Generates Credit Appraisal Memo (CAM) in .docx format using python-docx.
"""

import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.models import (
    Company, FinancialMetric, RiskFlag, ResearchFinding,
    PromoterAnalysis, RiskScore, SWOTAnalysis,
)

logger = logging.getLogger(__name__)
settings = get_settings()


class CAMReportService:
    """Generates Credit Appraisal Memo (CAM) reports."""

    def generate_report(self, company_id: str, db: Session) -> str:
        """
        Generate a full CAM report and return the file path.
        """
        company = db.query(Company).filter(Company.id == company_id).first()
        if not company:
            raise ValueError(f"Company not found: {company_id}")

        metrics = (
            db.query(FinancialMetric)
            .filter(FinancialMetric.company_id == company_id)
            .order_by(FinancialMetric.created_at.desc())
            .all()
        )
        risk_flags = (
            db.query(RiskFlag)
            .filter(RiskFlag.company_id == company_id)
            .all()
        )
        research = (
            db.query(ResearchFinding)
            .filter(ResearchFinding.company_id == company_id)
            .all()
        )
        promoters = (
            db.query(PromoterAnalysis)
            .filter(PromoterAnalysis.company_id == company_id)
            .all()
        )
        risk_score = (
            db.query(RiskScore)
            .filter(RiskScore.company_id == company_id)
            .order_by(RiskScore.created_at.desc())
            .first()
        )

        doc = DocxDocument()
        self._set_styles(doc)

        # Title Page
        self._add_title_page(doc, company)

        # Table of Contents
        self._add_section_header(doc, "Table of Contents")
        toc_items = [
            "1. Company Overview",
            "2. Financial Summary",
            "3. Financial Validation & Risk Flags",
            "4. External Intelligence Findings",
            "5. SWOT Analysis",
            "6. Promoter Risk Analysis",
            "7. Early Warning Signals",
            "8. Five Cs Credit Evaluation",
            "9. Risk Score & Explanation",
            "10. AI Reasoning Narrative",
            "11. Loan Recommendation",
        ]
        for item in toc_items:
            doc.add_paragraph(item, style="List Number")
        doc.add_page_break()

        # Section 1: Company Overview
        self._add_company_overview(doc, company)

        # Section 2: Financial Summary
        self._add_financial_summary(doc, metrics)

        # Section 3: Financial Validation
        self._add_risk_flags(doc, risk_flags)

        # Section 4: External Intelligence
        self._add_research_findings(doc, research)

        # Section 5: SWOT Analysis
        swot = (
            db.query(SWOTAnalysis)
            .filter(SWOTAnalysis.company_id == company_id)
            .order_by(SWOTAnalysis.created_at.desc())
            .first()
        )
        self._add_swot_analysis(doc, swot)

        # Section 6: Promoter Risk
        self._add_promoter_analysis(doc, promoters)

        # Section 7: Early Warning Signals
        self._add_early_warnings(doc, risk_flags)

        # Section 8: Five Cs
        self._add_five_cs(doc, risk_score)

        # Section 9: Risk Score
        self._add_risk_score(doc, risk_score)

        # Section 10: AI Reasoning Narrative
        self._add_reasoning_narrative(doc, risk_score)

        # Section 11: Recommendation
        self._add_recommendation(doc, risk_score, company)

        # Footer
        self._add_footer(doc)

        # Save
        os.makedirs(settings.reports_dir, exist_ok=True)
        safe_name = "".join(c for c in company.name if c.isalnum() or c in " -_")
        filename = f"CAM_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(settings.reports_dir, filename)
        doc.save(filepath)

        logger.info(f"CAM report generated: {filepath}")
        return filepath

    def _set_styles(self, doc: DocxDocument):
        """Configure document styles."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_title_page(self, doc: DocxDocument, company: Company):
        """Add the title page."""
        doc.add_paragraph("")
        doc.add_paragraph("")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("CREDIT APPRAISAL MEMORANDUM")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph("")

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(company.name)
        run.bold = True
        run.font.size = Pt(18)

        doc.add_paragraph("")

        details = doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details.add_run(f"Industry: {company.industry or 'N/A'}\n").font.size = Pt(12)
        details.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n").font.size = Pt(12)
        details.add_run(f"Loan Requested: ₹{company.loan_amount_requested:,.2f}\n" if company.loan_amount_requested else "").font.size = Pt(12)
        details.add_run("Prepared by: Intelli-Credit AI Platform\n").font.size = Pt(12)

        doc.add_paragraph("")
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run("CONFIDENTIAL")
        run.bold = True
        run.font.color.rgb = RGBColor(204, 0, 0)
        run.font.size = Pt(14)

        doc.add_page_break()

    def _add_section_header(self, doc: DocxDocument, text: str):
        """Add a section header."""
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

    def _add_company_overview(self, doc: DocxDocument, company: Company):
        """Section 1: Company Overview."""
        self._add_section_header(doc, "1. Company Overview")

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        fields = [
            ("Company Name", company.name),
            ("Industry", company.industry or "N/A"),
            ("CIN", company.cin or "N/A"),
            ("PAN", company.pan or "N/A"),
            ("GST Number", company.gst_number or "N/A"),
            ("Registered Address", company.registered_address or "N/A"),
            ("Incorporation Date", company.incorporation_date or "N/A"),
            ("Contact Email", company.contact_email or "N/A"),
            ("Contact Phone", company.contact_phone or "N/A"),
            ("Loan Amount Requested", f"₹{company.loan_amount_requested:,.2f}" if company.loan_amount_requested else "N/A"),
            ("Loan Purpose", company.loan_purpose or "N/A"),
        ]

        for label, value in fields:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)

        doc.add_paragraph("")

    def _add_financial_summary(self, doc: DocxDocument, metrics: list):
        """Section 2: Financial Summary."""
        self._add_section_header(doc, "2. Financial Summary")

        if not metrics:
            doc.add_paragraph("No financial data available.")
            return

        m = metrics[0]  # Most recent

        # Key metrics table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"

        fin_fields = [
            ("Fiscal Year", m.fiscal_year or "N/A"),
            ("Revenue", f"₹{m.revenue:,.2f}" if m.revenue else "N/A"),
            ("Net Profit", f"₹{m.net_profit:,.2f}" if m.net_profit else "N/A"),
            ("Gross Profit", f"₹{m.gross_profit:,.2f}" if m.gross_profit else "N/A"),
            ("EBITDA", f"₹{m.ebitda:,.2f}" if m.ebitda else "N/A"),
            ("Total Assets", f"₹{m.total_assets:,.2f}" if m.total_assets else "N/A"),
            ("Total Liabilities", f"₹{m.total_liabilities:,.2f}" if m.total_liabilities else "N/A"),
            ("Total Debt", f"₹{m.total_debt:,.2f}" if m.total_debt else "N/A"),
            ("Shareholders Equity", f"₹{m.shareholders_equity:,.2f}" if m.shareholders_equity else "N/A"),
        ]

        for label, value in fin_fields:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)

        doc.add_paragraph("")

        # Ratios table
        doc.add_heading("Financial Ratios", level=2)
        ratio_table = doc.add_table(rows=0, cols=2)
        ratio_table.style = "Light Grid Accent 1"

        ratios = [
            ("Debt Ratio", f"{m.debt_ratio:.2%}" if m.debt_ratio else "N/A"),
            ("Current Ratio", f"{m.current_ratio:.2f}" if m.current_ratio else "N/A"),
            ("Debt to Equity", f"{m.debt_to_equity:.2f}" if m.debt_to_equity else "N/A"),
            ("Interest Coverage", f"{m.interest_coverage:.2f}" if m.interest_coverage else "N/A"),
            ("Profit Margin", f"{m.profit_margin:.2%}" if m.profit_margin else "N/A"),
            ("Return on Assets", f"{m.return_on_assets:.2%}" if m.return_on_assets else "N/A"),
            ("Return on Equity", f"{m.return_on_equity:.2%}" if m.return_on_equity else "N/A"),
        ]

        for label, value in ratios:
            row = ratio_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)

        doc.add_paragraph("")

    def _add_risk_flags(self, doc: DocxDocument, flags: list):
        """Section 3: Financial Validation & Risk Flags."""
        self._add_section_header(doc, "3. Financial Validation & Risk Flags")

        if not flags:
            doc.add_paragraph("No financial discrepancies or risk flags detected.")
            return

        for flag in flags:
            p = doc.add_paragraph()
            severity_colors = {
                "low": RGBColor(0, 128, 0),
                "medium": RGBColor(255, 165, 0),
                "high": RGBColor(255, 69, 0),
                "critical": RGBColor(204, 0, 0),
            }
            color = severity_colors.get(flag.severity.value if flag.severity else "medium", RGBColor(0, 0, 0))

            run = p.add_run(f"[{flag.severity.value.upper() if flag.severity else 'UNKNOWN'}] ")
            run.bold = True
            run.font.color.rgb = color

            p.add_run(flag.description)

        doc.add_paragraph("")

    def _add_research_findings(self, doc: DocxDocument, findings: list):
        """Section 4: External Intelligence Findings."""
        self._add_section_header(doc, "4. External Intelligence Findings")

        if not findings:
            doc.add_paragraph("No external research findings available.")
            return

        categories = {}
        for f in findings:
            cat = f.category or "general"
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(f)

        for category, items in categories.items():
            doc.add_heading(category.replace("_", " ").title(), level=2)
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(f"{item.title or 'Finding'}: ")
                run.bold = True
                p.add_run(item.summary or "No summary available.")
                if item.sentiment:
                    sent_p = doc.add_paragraph(f"  Sentiment: {item.sentiment}")
                    sent_p.style = "List Bullet"

        doc.add_paragraph("")

    def _add_swot_analysis(self, doc: DocxDocument, swot):
        """Section 5: SWOT Analysis."""
        self._add_section_header(doc, "5. SWOT Analysis")

        if not swot:
            doc.add_paragraph("No SWOT analysis generated. Run SWOT generation first.")
            return

        if swot.summary:
            p = doc.add_paragraph()
            run = p.add_run("Executive Summary: ")
            run.bold = True
            p.add_run(swot.summary)
            doc.add_paragraph("")

        swot_sections = [
            ("Strengths", swot.strengths, RGBColor(34, 139, 34)),
            ("Weaknesses", swot.weaknesses, RGBColor(220, 53, 69)),
            ("Opportunities", swot.opportunities, RGBColor(0, 123, 255)),
            ("Threats", swot.threats, RGBColor(255, 140, 0)),
        ]

        for title, items, color in swot_sections:
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                run.font.color.rgb = color

            if items:
                for item in items:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(item)
            else:
                doc.add_paragraph("No data available for this section.")
            doc.add_paragraph("")

        if swot.data_sources:
            p = doc.add_paragraph()
            run = p.add_run("Data Sources Used: ")
            run.bold = True
            run.font.size = Pt(9)
            sources = []
            for key, val in swot.data_sources.items():
                sources.append(f"{key.replace('_', ' ').title()}: {val}")
            p.add_run(", ".join(sources)).font.size = Pt(9)

        doc.add_paragraph("")

    def _add_promoter_analysis(self, doc: DocxDocument, promoters: list):
        """Section 6: Promoter Risk Analysis."""
        self._add_section_header(doc, "6. Promoter Risk Analysis")

        if not promoters:
            doc.add_paragraph("No promoter analysis available.")
            return

        for p in promoters:
            doc.add_heading(p.promoter_name, level=2)

            table = doc.add_table(rows=0, cols=2)
            table.style = "Light Grid Accent 1"

            fields = [
                ("Designation", p.designation or "N/A"),
                ("Risk Level", p.risk_level.value.upper() if p.risk_level else "N/A"),
                ("Bankruptcy Flag", "YES" if p.bankruptcy_flag else "No"),
                ("Fraud Flag", "YES" if p.fraud_flag else "No"),
                ("Regulatory Violations", "YES" if p.regulatory_violation_flag else "No"),
            ]

            for label, value in fields:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)

            if p.background_summary:
                doc.add_paragraph(p.background_summary)
            if p.risk_summary:
                doc.add_paragraph(f"Risk Summary: {p.risk_summary}")

        doc.add_paragraph("")

    def _add_early_warnings(self, doc: DocxDocument, flags: list):
        """Section 7: Early Warning Signals."""
        self._add_section_header(doc, "7. Early Warning Signals")

        critical_flags = [f for f in flags if f.severity and f.severity.value in ("high", "critical")]

        if not critical_flags:
            doc.add_paragraph("No critical early warning signals detected.")
            return

        doc.add_paragraph(
            f"Total early warning signals: {len(critical_flags)}"
        )

        for flag in critical_flags:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{flag.severity.value.upper()}] ")
            run.bold = True
            run.font.color.rgb = RGBColor(204, 0, 0)
            p.add_run(flag.description)

        doc.add_paragraph("")

    def _add_five_cs(self, doc: DocxDocument, risk_score: Optional[RiskScore]):
        """Section 8: Five Cs Credit Evaluation."""
        self._add_section_header(doc, "8. Five Cs Credit Evaluation")

        if not risk_score or not risk_score.five_cs_evaluation:
            doc.add_paragraph("Five Cs evaluation not available.")
            return

        five_cs = risk_score.five_cs_evaluation
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        headers = table.rows[0].cells
        headers[0].text = "Category"
        headers[1].text = "Score"
        headers[2].text = "Max"
        headers[3].text = "Assessment"

        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for cs_name in ["character", "capacity", "capital", "collateral", "conditions"]:
            cs = five_cs.get(cs_name, {})
            row = table.add_row()
            row.cells[0].text = cs_name.title()
            row.cells[1].text = str(cs.get("score", "N/A"))
            row.cells[2].text = str(cs.get("max", 10))
            row.cells[3].text = cs.get("assessment", "N/A")

        doc.add_paragraph("")

        # Add reasoning bullets for each C if available
        for cs_name in ["character", "capacity", "capital", "collateral", "conditions"]:
            cs = five_cs.get(cs_name, {})
            reasoning = cs.get("reasoning", [])
            if reasoning:
                doc.add_heading(f"{cs_name.title()} — Reasoning", level=2)
                for reason in reasoning:
                    doc.add_paragraph(reason, style="List Bullet")
                doc.add_paragraph("")

    def _add_risk_score(self, doc: DocxDocument, risk_score: Optional[RiskScore]):
        """Section 9: Risk Score & Explanation."""
        self._add_section_header(doc, "9. Risk Score & Explanation")

        if not risk_score:
            doc.add_paragraph("Risk score not calculated.")
            return

        p = doc.add_paragraph()
        p.add_run(f"Probability of Default: ").bold = True
        p.add_run(f"{risk_score.probability_of_default:.2%}")

        p = doc.add_paragraph()
        p.add_run(f"Risk Level: ").bold = True
        p.add_run(risk_score.risk_level.value.upper() if risk_score.risk_level else "N/A")

        # Positive factors
        if risk_score.positive_factors:
            doc.add_heading("Positive Risk Factors", level=2)
            for factor in risk_score.positive_factors:
                doc.add_paragraph(factor, style="List Bullet")

        # Negative factors
        if risk_score.negative_factors:
            doc.add_heading("Negative Risk Factors", level=2)
            for factor in risk_score.negative_factors:
                doc.add_paragraph(factor, style="List Bullet")

        # Feature importance
        if risk_score.feature_importance:
            doc.add_heading("Feature Importance", level=2)
            fi_table = doc.add_table(rows=1, cols=2)
            fi_table.style = "Light Grid Accent 1"
            fi_table.rows[0].cells[0].text = "Feature"
            fi_table.rows[0].cells[1].text = "Importance"

            for feat, imp in list(risk_score.feature_importance.items())[:10]:
                row = fi_table.add_row()
                row.cells[0].text = feat.replace("_", " ").title()
                row.cells[1].text = f"{imp:.4f}"

        doc.add_paragraph("")

    def _add_reasoning_narrative(self, doc: DocxDocument, risk_score: Optional[RiskScore]):
        """Section 10: AI Reasoning Narrative — step-by-step explainability."""
        self._add_section_header(doc, "10. AI Reasoning Narrative")

        if not risk_score or not risk_score.reasoning_narrative:
            doc.add_paragraph(
                "No AI reasoning narrative available. Calculate risk score to generate."
            )
            return

        doc.add_paragraph(
            "The following narrative was generated by our AI engine to walk through the "
            "credit decision logic step by step. Each section traces the data sources, "
            "validates key metrics, and explains the reasoning behind the final recommendation."
        )
        doc.add_paragraph("")

        # Parse the narrative by sections (split on ## headers)
        narrative = risk_score.reasoning_narrative
        sections = narrative.split("\n## ")

        for i, section in enumerate(sections):
            if not section.strip():
                continue
            lines = section.strip().split("\n")
            # First line is the section title (may or may not have ## prefix)
            title = lines[0].lstrip("#").strip()
            body = "\n".join(lines[1:]).strip()

            if title:
                doc.add_heading(title, level=2)
            if body:
                # Handle bullet points
                for line in body.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("- ") or line.startswith("* "):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    elif line.startswith("**") and line.endswith("**"):
                        p = doc.add_paragraph()
                        run = p.add_run(line.strip("*"))
                        run.bold = True
                    else:
                        doc.add_paragraph(line)

        doc.add_paragraph("")

    def _add_recommendation(self, doc: DocxDocument, risk_score: Optional[RiskScore], company: Company):
        """Section 11: Final Loan Recommendation."""
        self._add_section_header(doc, "11. Final Loan Recommendation")

        if not risk_score:
            doc.add_paragraph("No risk assessment available for recommendation.")
            return

        # Decision
        decision_map = {
            "approve": ("APPROVE LOAN", RGBColor(0, 128, 0)),
            "approve_with_conditions": ("APPROVE WITH CONDITIONS", RGBColor(255, 165, 0)),
            "reject": ("REJECT LOAN", RGBColor(204, 0, 0)),
        }

        decision_val = risk_score.decision.value if risk_score.decision else "reject"
        label, color = decision_map.get(decision_val, ("UNKNOWN", RGBColor(0, 0, 0)))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n{label}\n")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = color

        doc.add_paragraph("")

        # Details table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"

        details = [
            ("Company", company.name),
            ("Loan Requested", f"₹{company.loan_amount_requested:,.2f}" if company.loan_amount_requested else "N/A"),
            ("Recommended Loan Limit", f"₹{risk_score.recommended_loan_limit:,.2f}" if risk_score.recommended_loan_limit else "N/A"),
            ("Suggested Interest Rate", f"{risk_score.suggested_interest_rate}%" if risk_score.suggested_interest_rate else "N/A"),
            ("Risk Score", f"{risk_score.probability_of_default:.2%}"),
            ("Risk Level", risk_score.risk_level.value.upper() if risk_score.risk_level else "N/A"),
        ]

        for label, value in details:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)

        doc.add_paragraph("")

        if risk_score.due_diligence_notes:
            doc.add_heading("Due Diligence Notes", level=2)
            doc.add_paragraph(risk_score.due_diligence_notes)

    def _add_footer(self, doc: DocxDocument):
        """Add report footer."""
        doc.add_paragraph("")
        doc.add_paragraph("---")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Generated by Intelli-Credit AI Platform | {datetime.now().strftime('%B %d, %Y %H:%M')} | CONFIDENTIAL"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)


cam_report_service = CAMReportService()
